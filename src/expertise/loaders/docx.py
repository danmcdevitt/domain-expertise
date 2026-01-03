"""DOCX document loader using python-docx."""

from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .base import Document, DocumentLoader


class DocxLoader(DocumentLoader):
    """Load Microsoft Word documents.
    
    Extracts text from paragraphs and tables.
    """
    
    supported_extensions = [".docx", ".doc"]
    
    def __init__(self, include_tables: bool = True, split_sections: bool = False):
        """Initialize DOCX loader.
        
        Args:
            include_tables: Include text from tables.
            split_sections: If True, return one Document per heading section.
        """
        self.include_tables = include_tables
        self.split_sections = split_sections
    
    def load(self, path: Path | str) -> list[Document]:
        """Load DOCX from file path."""
        path = Path(path)
        
        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")
        
        if path.suffix.lower() == ".doc":
            raise ValueError("Legacy .doc format not supported. Please convert to .docx")
        
        try:
            doc = DocxDocument(path)
        except PackageNotFoundError:
            raise ValueError(f"Could not open {path} - may be corrupted or password protected")
        
        if self.split_sections:
            return self._load_by_sections(doc, str(path))
        else:
            return self._load_as_single(doc, str(path))
    
    def load_text(self, text: str, source: str = "text") -> list[Document]:
        """DOCX can't be loaded from raw text."""
        raise NotImplementedError("DocxLoader cannot load from raw text")
    
    def _load_as_single(self, doc: DocxDocument, source: str) -> list[Document]:
        """Load entire document as single Document."""
        paragraphs = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Extract tables
        if self.include_tables:
            for table in doc.tables:
                table_text = self._extract_table(table)
                if table_text:
                    paragraphs.append(table_text)
        
        if not paragraphs:
            return []
        
        # Extract metadata from core properties
        metadata = {}
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        
        return [Document(
            content="\n\n".join(paragraphs),
            source=source,
            title=doc.core_properties.title,
            metadata=metadata,
        )]
    
    def _load_by_sections(self, doc: DocxDocument, source: str) -> list[Document]:
        """Split document by heading styles."""
        documents = []
        current_section = None
        current_content = []
        
        for para in doc.paragraphs:
            # Check if this is a heading
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_content:
                    documents.append(Document(
                        content="\n\n".join(current_content),
                        source=source,
                        section=current_section,
                    ))
                
                current_section = para.text.strip()
                current_content = []
            else:
                text = para.text.strip()
                if text:
                    current_content.append(text)
        
        # Save final section
        if current_content:
            documents.append(Document(
                content="\n\n".join(current_content),
                source=source,
                section=current_section,
            ))
        
        return documents
    
    def _extract_table(self, table) -> str:
        """Extract text from a table as formatted text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        return "\n".join(rows)
